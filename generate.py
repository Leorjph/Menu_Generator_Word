from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import StringIO
import re
from deep_translator import GoogleTranslator


textfile_path = 'menu.txt'
template_path = './Templates'
image_path = './Images'
images = {'caution' : '/caution.png'}
output_file_name = 'Menu.docx'
io_folder = "./"

TEXT_WIDTH = 6
IMG_WIDTH = 7.6-TEXT_WIDTH

FONT_PROFILES = {
    'iron_skillet' : {
        'name_fr' : {'name' : 'Archivo Black', 'size' : 21, 'bold' : True, 'color' : RGBColor(0, 0, 0)},
        'ingredients_fr' : {'name' : 'Times', 'size' : 12, 'bold' : True, 'color' : RGBColor(0, 0, 0)},
        'name_en' : {'name' : 'Archivo Black', 'size' : 14, 'bold' : True, 'color' : RGBColor(0, 0, 0)},
        'ingredients_en' : {'name' : 'Times', 'size' : 10, 'bold' : True, 'color' : RGBColor(0, 0, 0)},
        'allergens' : {'name' : 'Calibri', 'size' : 10, 'bold' : True, 'color' : RGBColor(255, 0, 0)}
    },
    'true_balance' : {
        'name_fr' : {'name' : 'Archivo Black', 'size' : 23, 'bold' : True, 'color' : RGBColor(1, 97, 153)},
        'ingredients_fr' : {'name' : 'Times', 'size' : 13, 'bold' : True, 'color' : RGBColor(1, 97, 153)},
        'name_en' : {'name' : 'Archivo Black', 'size' : 14, 'bold' : True, 'color' : RGBColor(1, 97, 153)},
        'ingredients_en' : {'name' : 'Times', 'size' : 10, 'bold' : True, 'color' : RGBColor(1, 97, 153)},
        'allergens' : {'name' : 'Calibri', 'size' : 10, 'bold' : True, 'color' : RGBColor(255, 0, 0)}
    },
    'motd' : {
        'name_fr' : {'name' : 'Poppins', 'size' : 15, 'bold' : True, 'color' : RGBColor(26, 78, 61)},
        'ingredients_fr' : {'name' : 'Poppins', 'size' : 11, 'bold' : False, 'color' : RGBColor(26, 78, 61)},
        'name_en' : {'name' : 'Poppins', 'size' : 11, 'bold' : True, 'color' : RGBColor(0, 0, 0)},
        'ingredients_en' : {'name' : 'Poppins', 'size' : 8, 'bold' : False, 'color' : RGBColor(0, 0, 0)},
        'allergens' : {'name' : 'Poppins', 'size' : 8, 'bold' : False, 'color' : RGBColor(255, 0, 0)}
    },
    'st' : {
        'name_fr' : {'name' : 'Ruda', 'size' : 17, 'bold' : True, 'color' : RGBColor(0, 0, 0)},
        'ingredients_fr' : {'name' : 'Ruda', 'size' : 10.5, 'bold' : False, 'color' : RGBColor(0, 0, 0)},
        'name_en' : {'name' : 'Poppins', 'size' : 10.5, 'bold' : True, 'color' : RGBColor(0, 0, 0)},
        'ingredients_en' : {'name' : 'Poppins', 'size' : 9, 'bold' : False, 'color' : RGBColor(0, 0, 0)},
        'allergens' : {'name' : 'Poppins', 'size' : 9, 'bold' : False, 'color' : RGBColor(255, 0, 0)}
    }
}


TEMPLATES = {
    'iron_skillet' : '/is_template.docx',
    'true_balance' : '/tb_template.docx',
    'motd' : '/motd_template.docx',
    'st' : '/st_template.docx'
}


ALIGNMENT = {
    'default' : {'text_width' : 6, 'img_width' : 1.6},
    'motd' : {'text_width' : 3.5, 'img_width' : 1},
    'st' : {'text_width' : 2.8, 'img_width' : 1.6}
}


TAGS = {
    'vegan' : '/vegan.png',
    'v' : '/vegan.png',
    'vegetarian' : '/vegetarian.png',
    'veg' : '/vegetarian.png',
    'eat well' : '/eat_well.png',
    'ew' : '/eat_well.png',
    'mb' : '/eat_well.png',
    'halal' : '/halal.png',
    'h' : '/halal.png',
    'low carbon' : '/low_carbon.png',
    'lc' : '/low_carbon.png',
    'fe' : '/low_carbon.png'
}


def parse_text(f = None):
    items = []
    if not f:
        f = open(io_folder + textfile_path)
    while True:
        line = f.readline()
        if line == '':
            break
        line = line.strip()
        match = re.fullmatch(r"\d+\. ?(.*)", line)
        if match:
            name = match.group(1).strip()
            ingredients = f.readline().strip()
            if ingredients == '-':
                    ingredients = None
            tags = f.readline().strip().lower()
            if tags == '-':
                    tags = None
            allergens = f.readline().strip()
            if allergens == '-':
                    allergens = None
            items.append({'name':name, 'ingredients':ingredients, 'tags':tags, 'allergens':allergens})
    f.close()
    return items


def set_table_indent(table, inches):
    tblPr = table._element.tblPr
    
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    dxa_value = int(inches * 1440)
    tblInd.set(qn('w:w'), str(dxa_value))
    tblInd.set(qn('w:type'), 'dxa')
    
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')
    table.autofit = False


def apply_font_profile(run, profile_name, station_name='iron_skillet', scale=1):
    profile = FONT_PROFILES[station_name].get(profile_name)
    if profile:
        run.font.name = profile['name']
        run.font.size = Pt(int(profile['size'] * scale))
        run.font.bold = profile['bold']
        run.font.color.rgb = profile['color']
        
        
def format_text_paragraphs(text_cell, add_para = True, spacing = Pt(10), space_before = Pt(5), space_after = Pt(5), alignment='left'):
    if add_para:
        text_para = text_cell.add_paragraph()
    else:
        text_para = text_cell.paragraphs[0]
    text_para.paragraph_format.line_spacing = spacing
    text_para.paragraph_format.space_before = space_before
    text_para.paragraph_format.space_after = space_after
    
    if alignment == 'center':
        text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return text_para


def styleDocument(doc, items, type='default', hasTags=True, scale=1):

    section = doc.sections[0]
    
    match type:
        case 'motd':
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(1.8)
            section.top_margin = Inches(0.85)
            section.bottom_margin = Inches(0.9)
        case _:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(0.85)
            section.bottom_margin = Inches(0.9)
    
    style = doc.styles['normal']

    match type:
        case 'motd':
            table = doc.add_table(rows=len(items)+1, cols=1)
            text_width = ALIGNMENT[type]['text_width'] + ALIGNMENT[type]['img_width']
            table.columns[0].width = Inches(text_width)
            set_table_indent(table, 0.5)
        case 'st':
            table = doc.add_table(rows=len(items)+1, cols=4)
            set_table_indent(table, -0.4)
            
            text_width = ALIGNMENT[type]['text_width']
            img_width = ALIGNMENT[type]['img_width']
            
            # First column 
            table.columns[0].width = Inches(text_width)
            for cell in table.columns[0].cells:
                cell.width = Inches(text_width)
                
            # Space between
            table.columns[1].width = Inches(img_width)
            for cell in table.columns[1].cells:
                cell.width = Inches(img_width)
                
            # Second column
            table.columns[2].width = Inches(text_width)
            for cell in table.columns[2].cells:
                cell.width = Inches(text_width)
                
            p = format_text_paragraphs(table.columns[0].cells[0], spacing = Pt(int(22 * scale)), add_para=True, alignment='center')
            run = p.add_run('\n\n\n\n\n')
        case _:
            table = doc.add_table(rows=len(items)+1, cols=2)
            set_table_indent(table, -0.5)

            text_width = ALIGNMENT[type]['text_width']
            img_width = ALIGNMENT[type]['img_width']
            # Column 1 for the text
            table.columns[0].width = Inches(text_width)
            for cell in table.columns[0].cells:
                cell.width = Inches(text_width)
            # Colume 2 for the label images
            table.columns[1].width = Inches(img_width)
            for cell in table.columns[1].cells:
                cell.width = Inches(img_width)
    return table


def create_doc(station_name, items, save=False, scale=1):
    doc = Document(template_path + TEMPLATES[station_name])
    
    def add_name_fr(item, text_cell):
        if item['name']:
            p = format_text_paragraphs(text_cell, spacing = Pt(int(22 * scale)), add_para=False, alignment=alignment)
            name_fr = GoogleTranslator(source="auto", target="fr").translate(item['name'])
            run = p.add_run(name_fr)
            apply_font_profile(run, 'name_fr', scale=scale, station_name=station_name)
    
    
    def add_name_en(item, text_cell):
        if item['name']:
            p = format_text_paragraphs(text_cell, spacing = Pt(int(15 * scale)), alignment=alignment)
            name_en = GoogleTranslator(source="auto", target="en").translate(item['name'])
            run = p.add_run(name_en)
            apply_font_profile(run, 'name_en', scale=scale, station_name=station_name)
    
    
    def add_ingredients_fr(item, text_cell):
        if item['ingredients']:
            p = format_text_paragraphs(text_cell, spacing = Pt(int(13 * scale)), alignment=alignment)
            ingredients_fr = GoogleTranslator(source="auto", target="fr").translate(item['ingredients'])
            run = p.add_run(ingredients_fr)
            apply_font_profile(run, 'ingredients_fr', scale=scale, station_name=station_name)
            
    
    def add_ingredients_en(item, text_cell):
        if item['ingredients']:
            p = format_text_paragraphs(text_cell, spacing = Pt(int(11 * scale)), alignment=alignment)
            ingredients_en = GoogleTranslator(source="auto", target="en").translate(item['ingredients'])
            run = p.add_run(ingredients_en)
            apply_font_profile(run, 'ingredients_en', scale=scale, station_name=station_name)
    
    
    def add_allergens(item, text_cell, hasTags):
        if item['allergens']:
            if hasTags:
                inner_table = text_cell.add_table(rows=1, cols=2)
                
                inner_table.columns[0].width = Inches(0.3)
                for cell in inner_table.columns[0].cells:
                    cell.width = Inches(0.3)
                inner_table.columns[1].width = Inches(TEXT_WIDTH-0.3)
                for cell in inner_table.columns[1].cells:
                    cell.width = Inches(TEXT_WIDTH-0.3)
                
                p = format_text_paragraphs(inner_table.cell(0, 0), add_para = False, alignment=alignment)
                run = p.add_run()
                run.add_picture(image_path + images['caution'], width=Inches(0.2 * scale))
                
                p = format_text_paragraphs(inner_table.cell(0, 1), add_para = False, spacing = Pt(int(10 * scale)), alignment=alignment)
            else:
                p = format_text_paragraphs(text_cell, spacing = Pt(int(11 * scale)), alignment=alignment)
            run = p.add_run()
            allergens = item['allergens']
            allergens_fr = GoogleTranslator(source="en", target="fr").translate(allergens)
            run = p.add_run(f'{allergens} / {allergens_fr}')
            apply_font_profile(run, 'allergens', scale=scale, station_name=station_name)
    

    hasTags = True
    if station_name == 'motd' or station_name == 'st':
        hasTags = False
        table = styleDocument(doc, items, type=station_name, hasTags=hasTags, scale=scale)
        alignment = 'center'
    else:
        table = styleDocument(doc, items)
        alignment = 'left'

    for i, item in enumerate(items):
        text_cell = table.cell(i+1, 0)
        if hasTags:
            image_cell = table.cell(i+1, 1)
        print(f"Generating item {i+1}")

        match station_name:
            case 'motd':
                add_name_fr(item, text_cell)
                add_name_en(item, text_cell)
                add_ingredients_fr(item, text_cell)
                add_ingredients_en(item, text_cell)
                add_allergens(item, text_cell, hasTags)
                
            case 'st':
                add_name_fr(item, text_cell)
                add_ingredients_fr(item, text_cell)
                add_name_en(item, text_cell)
                add_ingredients_en(item, text_cell)
                add_allergens(item, text_cell, hasTags)
                if item['tags']:
                    add_tags(text_cell, item['tags'], text_cell, scale=scale, alignment=alignment)
                text_cell = table.cell(i+1, 2)
                add_name_fr(item, text_cell)
                add_ingredients_fr(item, text_cell)
                add_name_en(item, text_cell)
                add_ingredients_en(item, text_cell)
                add_allergens(item, text_cell, hasTags)
                if item['tags']:
                    add_tags(text_cell, item['tags'], text_cell, scale=scale, alignment=alignment)
                    
            case _:
                add_name_fr(item, text_cell)
                add_ingredients_fr(item, text_cell)
                add_name_en(item, text_cell)
                add_ingredients_en(item, text_cell)
                add_allergens(item, text_cell, hasTags)
                if item['tags']:
                    add_tags(image_cell, item['tags'], text_cell, scale=scale, alignment=alignment)

    if save:
        doc.save(io_folder + output_file_name)
    return doc


def add_tags(cell, tags, text_cell, scale=1, alignment='left'):
    tag_list = tags.split(',')
    
    if alignment == 'left':
        rows = 1
        if len(tag_list) > 3:
            rows = 2
        cols = 3
        table = cell.add_table(rows=rows, cols=cols)
        for i in range(cols):
            table.columns[i].width = Inches(IMG_WIDTH/cols)
            for cell in table.columns[i].cells:
                cell.width = Inches(IMG_WIDTH/cols)
    elif alignment == 'center':
        p = format_text_paragraphs(cell, add_para = True, spacing = Inches(1.5 * scale), alignment=alignment)
        scale *= 1.5
    
    for i, tag in enumerate(tag_list[:9]):
        img_path = TAGS.get(tag.strip(), None)
        if img_path:
            try:
                if alignment == 'left':
                    p = format_text_paragraphs(table.cell(int(i/cols), i%cols), add_para = False, spacing = Inches(0.55 * scale))
                    run = p.add_run()
                else:
                    run = p.add_run(' ')
                run.add_picture(image_path + img_path, width=Inches(0.4 * scale))
            except FileNotFoundError:
                print(f"Error: File {img_path} not found")
                
    if alignment == 'left':
        cell.paragraphs[-1].add_run(' ').font.size = Pt(1)
        text_cell.paragraphs[-1].add_run(' ').font.size = Pt(1)


if __name__ == '__main__':
    items = parse_text()
    station_name = 'motd'
    create_doc(station_name, items, save=True)
    print(f"\nSuccessfully generated {output_file_name}\n")
